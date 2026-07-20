import os
import sys
import json
import calendar
import datetime
import logging
import tempfile
import threading
import tempfile as _tf
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from telegram import Update, ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes
from telegram.error import TimedOut, NetworkError

_log_dir = os.path.dirname(os.path.abspath(__file__))
_log_file = os.path.join(_log_dir, "bot.log")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(), logging.FileHandler(_log_file, encoding="utf-8")],
)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
if not BOT_TOKEN:
    logger.error("TELEGRAM_BOT_TOKEN environment variable is not set!")
    sys.exit(1)

DATA_DIR = "/app/data" if os.path.exists("/app/data") else os.path.dirname(__file__)
DATA_FILE = os.path.join(DATA_DIR, "headache_data.json")
_file_lock = threading.Lock()

DEFAULT_HOUR = 20
DEFAULT_MINUTE = 0

MONTHS_RU = [
    "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"
]

HARD_DATA = {
    "2026-02-01": False,
    "2026-02-02": False,
    "2026-02-03": False,
    "2026-02-04": False,
    "2026-02-05": False,
    "2026-02-06": False,
    "2026-02-07": False,
    "2026-02-08": False,
    "2026-02-09": False,
    "2026-02-10": False,
    "2026-02-11": False,
    "2026-02-12": False,
    "2026-02-13": False,
    "2026-02-14": False,
    "2026-02-15": False,
    "2026-02-16": False,
    "2026-02-17": False,
    "2026-02-18": False,
    "2026-02-19": True,
    "2026-02-20": False,
    "2026-02-21": False,
    "2026-02-22": True,
    "2026-02-23": True,
    "2026-02-24": True,
    "2026-02-25": False,
    "2026-02-26": True,
    "2026-02-27": False,
    "2026-02-28": False,
    "2026-03-01": False,
    "2026-03-02": False,
    "2026-03-03": False,
    "2026-03-04": False,
    "2026-03-05": True,
    "2026-03-06": False,
    "2026-03-07": True,
    "2026-03-08": False,
    "2026-03-09": False,
    "2026-03-10": False,
    "2026-03-11": False,
    "2026-03-12": True,
    "2026-03-13": False,
    "2026-03-14": False,
    "2026-03-15": True,
    "2026-03-16": False,
    "2026-03-17": True,
    "2026-03-18": True,
    "2026-03-19": False,
    "2026-03-20": True,
    "2026-03-21": False,
    "2026-03-22": True,
    "2026-03-23": True,
    "2026-03-24": False,
    "2026-03-25": False,
    "2026-03-26": True,
    "2026-03-27": False,
    "2026-03-28": True,
    "2026-03-29": False,
    "2026-03-30": False,
    "2026-03-31": True,
    "2026-04-01": True,
    "2026-04-02": True,
    "2026-04-03": True,
    "2026-04-04": True,
    "2026-04-05": False,
    "2026-04-06": False,
    "2026-04-07": True,
    "2026-04-08": True,
    "2026-04-09": False,
    "2026-04-10": False,
    "2026-04-11": True,
    "2026-04-12": False,
    "2026-04-13": False,
    "2026-04-14": True,
    "2026-04-15": True,
    "2026-04-16": False,
    "2026-04-17": False,
    "2026-04-18": True,
    "2026-04-19": True,
    "2026-04-20": False,
    "2026-04-21": False,
    "2026-04-22": False,
    "2026-04-23": False,
    "2026-04-24": True,
    "2026-04-25": True,
    "2026-04-26": False,
    "2026-04-27": False,
    "2026-04-28": False,
    "2026-04-29": False,
    "2026-04-30": False,
    "2026-05-01": True,
    "2026-05-02": False,
    "2026-05-03": False,
    "2026-05-04": True,
    "2026-05-05": False,
    "2026-05-06": False,
    "2026-05-07": True,
    "2026-05-08": False,
    "2026-05-09": True,
    "2026-05-10": False,
    "2026-05-11": False,
    "2026-05-12": False,
    "2026-05-13": False,
    "2026-05-14": False,
    "2026-05-15": False,
    "2026-05-16": False,
    "2026-05-17": True,
    "2026-05-18": False,
    "2026-05-19": False,
    "2026-05-20": True,
    "2026-05-21": True,
    "2026-05-22": True,
    "2026-05-23": False,
    "2026-05-24": True,
    "2026-05-25": False,
    "2026-05-26": True,
    "2026-05-27": False,
    "2026-05-28": True,
    "2026-05-29": True,
    "2026-05-30": False,
    "2026-05-31": False,
    "2026-06-01": False,
    "2026-06-02": False,
    "2026-06-03": False,
    "2026-06-04": True,
    "2026-06-05": True,
    "2026-06-06": True,
    "2026-06-07": False,
    "2026-06-08": False,
    "2026-06-09": True,
    "2026-06-10": False,
    "2026-06-11": False,
    "2026-06-12": False,
    "2026-06-13": False,
    "2026-06-14": True,
    "2026-06-15": True,
    "2026-06-16": False,
    "2026-06-17": False,
    "2026-06-18": True,
    "2026-06-19": False,
    "2026-06-20": False,
    "2026-06-21": False,
    "2026-06-22": False,
    "2026-06-23": True,
    "2026-06-24": True,
    "2026-07-01": False,
    "2026-07-02": False,
    "2026-07-03": True,
    "2026-07-04": True,
    "2026-07-05": False,
    "2026-07-06": True,
    "2026-07-07": False,
    "2026-07-08": False,
    "2026-07-09": True,
    "2026-07-10": False,
    "2026-07-11": True,
    "2026-07-12": True,
    "2026-07-13": False,
    "2026-07-14": False,
    "2026-07-15": False,
    "2026-07-16": False,
    "2026-07-17": False,
    "2026-07-18": False,
    "2026-07-19": False,
    "2026-07-20": False,
    "2026-07-21": False,
    "2026-07-22": False,
    "2026-07-23": False,
    "2026-07-24": False,
    "2026-07-25": False,
    "2026-07-26": False,
    "2026-07-27": False,
    "2026-07-28": False,
    "2026-07-29": False,
    "2026-07-30": False,
    "2026-07-31": False,
}


def _atomic_save(filepath, data):
    with _file_lock:
        dir_name = os.path.dirname(filepath) or "."
        fd, tmp_path = tempfile.mkstemp(dir=dir_name, suffix=".tmp")
        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            os.replace(tmp_path, filepath)
        except Exception:
            try:
                os.remove(tmp_path)
            except OSError:
                pass
            raise


def load_data():
    with _file_lock:
        logger.info("load_data: DATA_FILE=%s, exists=%s", DATA_FILE, os.path.exists(DATA_FILE))
        if not os.path.exists(DATA_FILE):
            return {}
        try:
            with open(DATA_FILE, "r", encoding="utf-8") as f:
                result = json.load(f)
                logger.info("load_data: loaded %d users", len(result))
                return result
        except (json.JSONDecodeError, OSError) as e:
            logger.error("load_data: error reading %s: %s", DATA_FILE, e)
            return {}


def save_data(data):
    logger.info("save_data: DATA_FILE=%s, users=%d", DATA_FILE, len(data))
    _atomic_save(DATA_FILE, data)


def get_user_data(user_id):
    data = load_data()
    uid = str(user_id)
    if uid not in data:
        data[uid] = {"answers": HARD_DATA.copy(), "hour": DEFAULT_HOUR, "minute": DEFAULT_MINUTE}
        save_data(data)
    if "hour" not in data[uid]:
        data[uid]["hour"] = DEFAULT_HOUR
    if "minute" not in data[uid]:
        data[uid]["minute"] = DEFAULT_MINUTE
    answers = data[uid].setdefault("answers", {})
    for k, v in HARD_DATA.items():
        if k not in answers:
            answers[k] = v
    return data[uid]


def save_user_data(user_id, user_data):
    clean = dict(user_data)
    clean_answers = {k: v for k, v in clean.get("answers", {}).items() if k not in HARD_DATA}
    clean["answers"] = clean_answers
    logger.info("save_user_data: keys=%s, notes_count=%s", list(clean.keys()), len(clean.get("notes", {})))
    data = load_data()
    data[str(user_id)] = clean
    save_data(data)


def get_user_time(user_id):
    user_data = get_user_data(user_id)
    return user_data.get("hour", DEFAULT_HOUR), user_data.get("minute", DEFAULT_MINUTE)


def get_main_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton("📋 Календарь"), KeyboardButton("⚙️ Настройки")],
            [KeyboardButton("📊 Статистика"), KeyboardButton("ℹ️ Помощь")],
        ],
        resize_keyboard=True,
    )


async def _safe_reply(message, text, **kwargs):
    for attempt in range(3):
        try:
            return await message.reply_text(text, **kwargs)
        except TimedOut:
            logger.warning("TimedOut on attempt %d, retrying...", attempt + 1)
            if attempt == 2:
                raise
        except NetworkError as e:
            logger.warning("NetworkError on attempt %d: %s", attempt + 1, e)
            if attempt == 2:
                raise


async def _safe_edit(query, text, **kwargs):
    try:
        return await query.edit_message_text(text, **kwargs)
    except Exception as e:
        logger.warning("edit_message_text failed: %s", e)
        return None


def get_settings_keyboard(hour, minute):
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton("◀️", callback_data="set_hour_prev"),
            InlineKeyboardButton(f"Час: {hour:02d}", callback_data="set_ignore"),
            InlineKeyboardButton("▶️", callback_data="set_hour_next"),
        ],
        [
            InlineKeyboardButton("◀️", callback_data="set_min_prev"),
            InlineKeyboardButton(f"Мин: {minute:02d}", callback_data="set_ignore"),
            InlineKeyboardButton("▶️", callback_data="set_min_next"),
        ],
        [InlineKeyboardButton("✅ Сохранить", callback_data="set_save")],
        [InlineKeyboardButton("⬅️ В меню", callback_data="set_back")],
    ])


async def reschedule_user_job(context, user_id, hour, minute):
    job_queue = context.job_queue
    job_name = f"daily_{user_id}"
    current_jobs = job_queue.get_jobs_by_name(job_name)
    for job in current_jobs:
        job.schedule_removal()
    utc_hour = (hour - 3) % 24
    job_queue.run_daily(
        send_daily_question,
        time=datetime.time(hour=utc_hour, minute=minute, second=0),
        chat_id=user_id,
        name=job_name,
    )
    logger.info("Rescheduled daily question for user %s at %02d:%02d MSK (UTC %02d:%02d)", user_id, hour, minute, utc_hour, minute)


def get_calendar_keyboard(user_id, year, month):
    user_data = get_user_data(user_id)
    answers = dict(HARD_DATA)
    data_dict = load_data()
    uid = str(user_id)
    file_answers = data_dict.get(uid, {}).get("answers", {})
    answers.update(file_answers)
    notes = user_data.get("notes", {})
    logger.info("Calendar: uid=%s, file_answers=%s, notes=%s", uid, file_answers, list(notes.keys()) if notes else [])

    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1

    buttons = []
    buttons.append([
        InlineKeyboardButton(f"◀ {MONTHS_RU[prev_month-1]}", callback_data=f"cal_prev_{year}_{month}"),
        InlineKeyboardButton(f"{MONTHS_RU[month-1]} {year}", callback_data="cal_ignore"),
        InlineKeyboardButton(f"{MONTHS_RU[next_month-1]} ▶", callback_data=f"cal_next_{year}_{month}"),
    ])
    buttons.append([
        InlineKeyboardButton("Пн", callback_data="cal_ignore"),
        InlineKeyboardButton("Вт", callback_data="cal_ignore"),
        InlineKeyboardButton("Ср", callback_data="cal_ignore"),
        InlineKeyboardButton("Чт", callback_data="cal_ignore"),
        InlineKeyboardButton("Пт", callback_data="cal_ignore"),
        InlineKeyboardButton("Сб", callback_data="cal_ignore"),
        InlineKeyboardButton("Вс", callback_data="cal_ignore"),
    ])

    cal = calendar.monthcalendar(year, month)
    pain_count = 0
    pain_dates = []
    for week in cal:
        row = []
        for day in week:
            if day == 0:
                row.append(InlineKeyboardButton(" ", callback_data="cal_ignore"))
            else:
                date_str = f"{year}-{month:02d}-{day:02d}"
                if date_str in answers:
                    if answers[date_str]:
                        marker = "🔺"
                        pain_count += 1
                        pain_dates.append(day)
                    else:
                        marker = "✓"
                    row.append(InlineKeyboardButton(f"{day}{marker}", callback_data=f"cal_day_{year}_{month}_{day}"))
                else:
                    row.append(InlineKeyboardButton(f"{day}", callback_data=f"cal_day_{year}_{month}_{day}"))
        buttons.append(row)

    max_streak = 0
    current_streak = 0
    for day in range(1, 32):
        if day in pain_dates:
            current_streak += 1
            max_streak = max(max_streak, current_streak)
        else:
            current_streak = 0

    total_days = calendar.monthrange(year, month)[1]
    pain_pct = round(pain_count / total_days * 100) if total_days > 0 else 0

    note_count = sum(1 for d in notes if d.startswith(f"{year}-{month:02d}"))
    buttons.append([
        InlineKeyboardButton("📝 Отчёт", callback_data=f"cal_report_{year}_{month}"),
    ])
    buttons.append([InlineKeyboardButton("⬅️ В меню", callback_data="cal_back")])

    header = (
        f"{MONTHS_RU[month-1]} {year}\n"
        f"🔺 {pain_count} дн. болела\n"
        f"🔥 До {max_streak} дн. подряд\n"
        f"📊 {pain_pct}% болезненных дней\n"
        f"📝 {note_count} заметок"
    )
    return InlineKeyboardMarkup(buttons), header


def get_report_calendar_keyboard(user_id, year, month, selected_start):
    user_data = get_user_data(user_id)
    answers = dict(HARD_DATA)
    data_dict = load_data()
    uid = str(user_id)
    file_answers = data_dict.get(uid, {}).get("answers", {})
    answers.update(file_answers)
    notes = user_data.get("notes", {})

    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1

    buttons = []
    buttons.append([
        InlineKeyboardButton(f"◀ {MONTHS_RU[prev_month-1]}", callback_data=f"cal_prev_{year}_{month}"),
        InlineKeyboardButton(f"{MONTHS_RU[month-1]} {year}", callback_data="cal_ignore"),
        InlineKeyboardButton(f"{MONTHS_RU[next_month-1]} ▶", callback_data=f"cal_next_{year}_{month}"),
    ])
    buttons.append([
        InlineKeyboardButton("Пн", callback_data="cal_ignore"),
        InlineKeyboardButton("Вт", callback_data="cal_ignore"),
        InlineKeyboardButton("Ср", callback_data="cal_ignore"),
        InlineKeyboardButton("Чт", callback_data="cal_ignore"),
        InlineKeyboardButton("Пт", callback_data="cal_ignore"),
        InlineKeyboardButton("Сб", callback_data="cal_ignore"),
        InlineKeyboardButton("Вс", callback_data="cal_ignore"),
    ])

    cal = calendar.monthcalendar(year, month)
    for week in cal:
        row = []
        for day in week:
            if day == 0:
                row.append(InlineKeyboardButton(" ", callback_data="cal_ignore"))
            else:
                date_str = f"{year}-{month:02d}-{day:02d}"
                note_marker = "📝" if date_str in notes else ""
                if selected_start and date_str == selected_start:
                    row.append(InlineKeyboardButton(f"🟢{day}{note_marker}", callback_data=f"cal_day_{year}_{month}_{day}"))
                else:
                    row.append(InlineKeyboardButton(f"{day}{note_marker}", callback_data=f"cal_day_{year}_{month}_{day}"))
        buttons.append(row)

    buttons.append([InlineKeyboardButton("⬅️ В меню", callback_data="cal_back")])
    header = f"📝 Выбор периода для отчёта\n{MONTHS_RU[month-1]} {year}"
    return InlineKeyboardMarkup(buttons), header


def generate_report(user_id, date_start, date_end):
    data_dict = load_data()
    uid = str(user_id)
    user_info = data_dict.get(uid, {})
    logger.info("Report: uid=%s, keys=%s", uid, list(user_info.keys()))
    notes = user_info.get("notes", {})
    logger.info("Report: notes keys=%s", list(notes.keys()) if notes else "EMPTY")
    answers = dict(HARD_DATA)
    file_answers = user_info.get("answers", {})
    answers.update(file_answers)

    d_start = datetime.datetime.strptime(date_start, "%Y-%m-%d").date()
    d_end = datetime.datetime.strptime(date_end, "%Y-%m-%d").date()

    report_notes = {}
    for date_str, note_text in notes.items():
        try:
            d = datetime.datetime.strptime(date_str, "%Y-%m-%d").date()
            if d_start <= d <= d_end:
                report_notes[date_str] = note_text
        except Exception:
            continue

    if not report_notes:
        return None

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    title = doc.add_heading('Отчёт по головной боли', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Период: {d_start.strftime('%d.%m.%Y')} — {d_end.strftime('%d.%m.%Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    pain_count = 0
    total_days = 0
    current = d_start
    while current <= d_end:
        date_str = current.strftime("%Y-%m-%d")
        if date_str in answers:
            total_days += 1
            if answers[date_str]:
                pain_count += 1
        current += datetime.timedelta(days=1)

    pain_pct = round(pain_count / total_days * 100) if total_days > 0 else 0

    doc.add_heading('Статистика', level=1)
    doc.add_paragraph(f"Всего дней в периоде: {total_days}")
    doc.add_paragraph(f"Дней с головной болью: {pain_count} ({pain_pct}%)")
    doc.add_paragraph(f"Дней без боли: {total_days - pain_count}")
    doc.add_paragraph()

    doc.add_heading('Заметки', level=1)

    sorted_dates = sorted(report_notes.keys())
    for date_str in sorted_dates:
        try:
            d = datetime.datetime.strptime(date_str, "%Y-%m-%d").date()
            date_display = d.strftime("%d.%m.%Y")
        except Exception:
            date_display = date_str

        pain_status = ""
        if date_str in answers:
            pain_status = " — болела голова" if answers[date_str] else " — не болела"

        doc.add_heading(f"{date_display}{pain_status}", level=2)
        doc.add_paragraph(report_notes[date_str])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Создано: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

    tmp = _tf.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    hour, minute = get_user_time(user_id)

    data = load_data()
    uid = str(user_id)
    if uid not in data:
        data[uid] = {"answers": {}, "hour": DEFAULT_HOUR, "minute": DEFAULT_MINUTE}
        save_data(data)
    user_data = data[uid]

    await reschedule_user_job(context, user_id, hour, minute)

    await update.message.reply_text(
        "Привет! Я дневник головной боли.\n"
        "Каждый день в {hour}:{minute:02d} я буду спрашивать, болела ли голова.\n\n"
        "В меню «⚙️ Настройки» можно изменить время.".format(
            hour=hour, minute=minute
        ),
        reply_markup=get_main_keyboard(),
    )


async def cmd_setpain(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    args = context.args
    if not args or len(args) < 2:
        await update.message.reply_text(
            "Использование: /setpain ДД.ММ.ГГГГ да/нет\n"
            "Пример: /setpain 13.07.2026 да"
        )
        return

    date_str_raw = args[0]
    pain_str = args[1].lower()

    try:
        date_obj = datetime.datetime.strptime(date_str_raw, "%d.%m.%Y").date()
    except ValueError:
        await update.message.reply_text("Неверный формат даты. Используй ДД.ММ.ГГГГ")
        return

    has_pain = pain_str in ("да", "yes", "1", "true")
    date_key = date_obj.strftime("%Y-%m-%d")

    data_dict = load_data()
    uid = str(user_id)
    if uid not in data_dict:
        data_dict[uid] = {"answers": {}, "hour": DEFAULT_HOUR, "minute": DEFAULT_MINUTE}
    if "answers" not in data_dict[uid]:
        data_dict[uid]["answers"] = {}
    data_dict[uid]["answers"][date_key] = has_pain
    save_data(data_dict)

    status = "болела голова" if has_pain else "не болела"
    await update.message.reply_text(
        f"✅ {date_str_raw} — {status}",
        reply_markup=get_main_keyboard(),
    )


async def cmd_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    hour, minute = get_user_time(user_id)
    await update.message.reply_text(
        "Как пользоваться:\n\n"
        "Каждый день в {hour}:{minute:02d} я пришлю вопрос: болела ли голова.\n"
        "Нажми «Да, болела» или «Нет, не болела».\n"
        "После ответа можно будет добавить заметку.\n\n"
        "Меню:\n"
        "📋 Календарь — посмотреть дни за текущий месяц\n"
        "⚙️ Настройки — изменить время вопроса\n"
        "📊 Статистика — общая статистика\n"
        "ℹ️ Помощь — эта справка".format(
            hour=hour, minute=minute
        ),
        reply_markup=get_main_keyboard(),
    )


async def send_daily_question(context: ContextTypes.DEFAULT_TYPE):
    chat_id = context.job.chat_id
    yesterday = (datetime.date.today() - datetime.timedelta(days=1)).isoformat()

    data = load_data()
    uid = str(chat_id)
    raw_answers = data.get(uid, {}).get("answers", {})

    if yesterday in raw_answers:
        return

    keyboard = InlineKeyboardMarkup([
        [
            InlineKeyboardButton("Да 🙄", callback_data=f"pain_yes_{yesterday}"),
            InlineKeyboardButton("Нет 🙂", callback_data=f"pain_no_{yesterday}"),
        ],
    ])

    try:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Привет. У вас вчера болела голова?",
            reply_markup=keyboard,
        )
    except Exception as e:
        logger.error("Failed to send daily question to %s: %s", chat_id, e)


async def handle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    data = query.data

    if data.startswith("pain_"):
        parts = data.split("_")
        has_pain = parts[1] == "yes"
        date_str = parts[2]
        logger.info("Pain callback received: date=%s, has_pain=%s, uid=%s", date_str, has_pain, user_id)

        data_dict = load_data()
        uid = str(user_id)
        if uid not in data_dict:
            data_dict[uid] = {"answers": {}, "hour": DEFAULT_HOUR, "minute": DEFAULT_MINUTE}
        user_data = data_dict[uid]
        if "answers" not in user_data:
            user_data["answers"] = {}
        user_data["answers"][date_str] = has_pain
        save_data(data_dict)
        logger.info("Saved pain answer: date=%s, has_pain=%s, uid=%s", date_str, has_pain, uid)

        emoji = "😣" if has_pain else "😊"
        text = "{emoji} Записал: {date} — {status}".format(
            emoji=emoji,
            date=datetime.datetime.strptime(date_str, "%Y-%m-%d").strftime("%d.%m.%Y"),
            status="голова болела" if has_pain else "голова не болела",
        )
        await _safe_edit(query, text)

        note_keyboard = InlineKeyboardMarkup([
            [
                InlineKeyboardButton("Да", callback_data=f"note_yes_{date_str}"),
                InlineKeyboardButton("Нет", callback_data=f"note_no_{date_str}"),
            ],
        ])
        await _safe_reply(query.message, "Хотите сделать заметку?", reply_markup=note_keyboard)
        return

    if data.startswith("note_"):
        parts = data.split("_")
        action = parts[1]
        date_str = parts[2]

        if action == "no":
            await _safe_edit(query, "Ок")
            await _safe_reply(query.message, "Выбери действие:", reply_markup=get_main_keyboard())
            return

        if action == "yes":
            data_dict = load_data()
            uid = str(user_id)
            if uid not in data_dict:
                data_dict[uid] = {"answers": {}, "hour": DEFAULT_HOUR, "minute": DEFAULT_MINUTE}
            data_dict[uid]["note_pending"] = date_str
            save_data(data_dict)
            logger.info("Set note_pending=%s for uid=%s", date_str, uid)
            await _safe_edit(query, "Напиши заметку:")
            return

        return

    if data.startswith("cal_prev_") or data.startswith("cal_next_"):
        parts = data.split("_")
        year = int(parts[2])
        month = int(parts[3])
        if data.startswith("cal_prev_"):
            month -= 1
            if month < 1:
                month = 12
                year -= 1
        else:
            month += 1
            if month > 12:
                month = 1
                year += 1
        keyboard, header = get_calendar_keyboard(user_id, year, month)
        await _safe_edit(query, header, reply_markup=keyboard)
        return

    if data == "cal_back":
        await _safe_edit(query, "Выбери действие:")
        await _safe_reply(query.message, "Выбери действие:", reply_markup=get_main_keyboard())
        return

    if data == "cal_ignore":
        return

    if data.startswith("cal_report_"):
        parts = data.split("_")
        year = int(parts[2])
        month = int(parts[3])
        data_dict = load_data()
        uid = str(user_id)
        data_dict.setdefault(uid, {})
        data_dict[uid]["report_start"] = None
        data_dict[uid]["report_year"] = year
        data_dict[uid]["report_month"] = month
        save_data(data_dict)
        keyboard, header = get_report_calendar_keyboard(user_id, year, month, None)
        await _safe_edit(query, f"📝 Выбери начальную дату:\n{header}", reply_markup=keyboard)
        return

    if data.startswith("cal_day_"):
        parts = data.split("_")
        year = int(parts[2])
        month = int(parts[3])
        day = int(parts[4])
        data_dict = load_data()
        uid = str(user_id)
        user_info = data_dict.get(uid, {})
        report_start = user_info.get("report_start")
        report_year = user_info.get("report_year")
        report_month = user_info.get("report_month")

        if report_start is not None and report_year == year and report_month == month:
            date_start = report_start
            date_end = f"{year}-{month:02d}-{day:02d}"
            if date_start > date_end:
                date_start, date_end = date_end, date_start
            data_dict[uid].pop("report_start", None)
            data_dict[uid].pop("report_year", None)
            data_dict[uid].pop("report_month", None)
            save_data(data_dict)
            docx_path = generate_report(user_id, date_start, date_end)
            if docx_path:
                await _safe_edit(query, f"📝 Отчёт за период:\n{date_start} — {date_end}")
                with open(docx_path, "rb") as f:
                    await query.message.reply_document(
                        document=f,
                        filename="Отчёт.docx",
                        caption=f"Отчёт за {date_start} — {date_end}",
                    )
                os.remove(docx_path)
            else:
                await _safe_edit(query, "Нет заметок за выбранный период.")
            await _safe_reply(query.message, "Выбери действие:", reply_markup=get_main_keyboard())
        else:
            data_dict[uid]["report_start"] = f"{year}-{month:02d}-{day:02d}"
            data_dict[uid]["report_year"] = year
            data_dict[uid]["report_month"] = month
            save_data(data_dict)
            keyboard, header = get_report_calendar_keyboard(user_id, year, month, f"{year}-{month:02d}-{day:02d}")
            await _safe_edit(query, f"📝 Начало: {day}.{month:02d}.{year}\nВыбери конечную дату:\n{header}", reply_markup=keyboard)
        return

    if data.startswith("set_"):
        user_data = get_user_data(user_id)
        hour = user_data.get("hour", DEFAULT_HOUR)
        minute = user_data.get("minute", DEFAULT_MINUTE)

        if data == "set_hour_next":
            hour = (hour + 1) % 24
        elif data == "set_hour_prev":
            hour = (hour - 1) % 24
        elif data == "set_min_next":
            minute = (minute + 15) % 60
        elif data == "set_min_prev":
            minute = (minute - 15) % 60
        elif data == "set_save":
            save_user_data(user_id, user_data)
            await reschedule_user_job(context, user_id, hour, minute)
            try:
                await query.message.delete()
            except Exception:
                pass
            await _safe_reply(
                query.message,
                "✅ Время сохранено: {hour}:{minute:02d}".format(
                    hour=hour, minute=minute
                ),
                reply_markup=get_main_keyboard(),
            )
            return
        elif data == "set_back":
            try:
                await query.message.delete()
            except Exception:
                pass
            await _safe_reply(query.message, "Выбери действие:", reply_markup=get_main_keyboard())
            return
        elif data == "set_ignore":
            return

        user_data["hour"] = hour
        user_data["minute"] = minute
        save_user_data(user_id, user_data)

        keyboard = get_settings_keyboard(hour, minute)
        await _safe_edit(
            query,
            "Настройки времени вопроса:\nТекущее время: {hour}:{minute:02d}".format(
                hour=hour, minute=minute
            ),
            reply_markup=keyboard,
        )
        return


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    text = update.message.text.strip()

    user_data = get_user_data(user_id)
    note_pending = user_data.get("note_pending")
    logger.info("handle_message: uid=%s, note_pending=%s, keys=%s", user_id, note_pending, list(user_data.keys()))
    if note_pending:
        if "notes" not in user_data:
            user_data["notes"] = {}
        user_data["notes"][note_pending] = text
        user_data.pop("note_pending", None)
        save_user_data(user_id, user_data)
        logger.info("Saved note for %s: %s", note_pending, text[:50])
        await update.message.reply_text(
            "✅ Заметка к {date} сохранена.".format(
                date=datetime.datetime.strptime(note_pending, "%Y-%m-%d").strftime("%d.%m.%Y")
            ),
            reply_markup=get_main_keyboard(),
        )
        return

    if text == "📋 Календарь":
        now = datetime.date.today()
        keyboard, header = get_calendar_keyboard(user_id, now.year, now.month)
        await update.message.reply_text(header, reply_markup=keyboard)
        return

    if text == "⚙️ Настройки":
        hour, minute = get_user_time(user_id)
        keyboard = get_settings_keyboard(hour, minute)
        await update.message.reply_text(
            "Настройки времени вопроса:\nТекущее время: {hour}:{minute:02d}".format(
                hour=hour, minute=minute
            ),
            reply_markup=keyboard,
        )
        return

    if text == "📊 Статистика":
        user_data = get_user_data(user_id)
        answers = dict(HARD_DATA)
        data_dict = load_data()
        uid = str(user_id)
        file_answers = data_dict.get(uid, {}).get("answers", {})
        answers.update(file_answers)
        total = len(answers)
        pain_days = sum(1 for v in answers.values() if v)
        no_pain_days = total - pain_days

        if total == 0:
            await update.message.reply_text("Пока нет записей.", reply_markup=get_main_keyboard())
            return

        pain_pct = round(pain_days / total * 100)

        monthly = {}
        for date_str, has_pain in answers.items():
            try:
                ym = date_str[:7]
            except Exception:
                continue
            if ym not in monthly:
                monthly[ym] = {"pain": 0, "total": 0}
            monthly[ym]["total"] += 1
            if has_pain:
                monthly[ym]["pain"] += 1

        month_lines = []
        for ym in sorted(monthly.keys()):
            m = monthly[ym]
            year, month = ym.split("-")
            month_name = MONTHS_RU[int(month) - 1]
            pct = round(m["pain"] / m["total"] * 100) if m["total"] > 0 else 0
            month_lines.append(
                "  {name} {year}: {pain}/{total} ({pct}%)".format(
                    name=month_name, year=year, pain=m["pain"], total=m["total"], pct=pct
                )
            )

        msg = (
            "📊 Статистика:\n\n"
            "Всего дней: {total}\n"
            "😣 Болела голова: {pain} ({pct}%)\n"
            "😊 Не болела: {no_pain}\n\n"
            "Головные боли в месяц:\n{months}".format(
                total=total, pain=pain_days, pct=pain_pct, no_pain=no_pain_days,
                months="\n".join(month_lines)
            )
        )
        await update.message.reply_text(msg, reply_markup=get_main_keyboard())
        return

    if text == "ℹ️ Помощь":
        await cmd_help(update, context)
        return

    await update.message.reply_text("Выбери действие:", reply_markup=get_main_keyboard())


async def error_handler(update, context):
    logger.error("Exception while handling an update: %s", context.error, exc_info=context.error)


async def post_init(application):
    data = load_data()
    for uid, udata in data.items():
        hour = udata.get("hour", DEFAULT_HOUR)
        minute = udata.get("minute", DEFAULT_MINUTE)
        utc_hour = (hour - 3) % 24
        application.job_queue.run_daily(
            send_daily_question,
            time=datetime.time(hour=utc_hour, minute=minute, second=0),
            chat_id=int(uid),
            name=f"daily_{uid}",
        )
        logger.info("Scheduled daily question for user %s at %02d:%02d MSK (UTC %02d:%02d)", uid, hour, minute, utc_hour, minute)


def main():
    logger.info("Bot starting... v4 May=%d", len(HARD_DATA))
    try:
        app = (
            ApplicationBuilder()
            .token(BOT_TOKEN)
            .connect_timeout(30)
            .read_timeout(30)
            .write_timeout(30)
            .post_init(post_init)
            .build()
        )
        app.add_handler(CommandHandler("start", cmd_start))
        app.add_handler(CommandHandler("help", cmd_help))
        app.add_handler(CommandHandler("setpain", cmd_setpain))
        app.add_handler(CallbackQueryHandler(handle_callback))
        app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
        app.add_error_handler(error_handler)
        logger.info("Bot is running!")
        app.run_polling(drop_pending_updates=True)
    except Exception as e:
        logger.exception("Bot crashed: %s", e)
        raise


if __name__ == "__main__":
    main()
